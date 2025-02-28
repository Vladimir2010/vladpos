import csv
import os
import sqlite3
import platform
import subprocess
import sys
import psycopg2
from docx import Document
from PyQt6.QtCore import Qt,QDate
from PyQt6.QtGui import QAction, QIntValidator, QDoubleValidator, QPainter, QPixmap, QIcon
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog  # За работа с печат
from PyQt6.QtWidgets import QApplication, QMainWindow, QTableWidget, QTableWidgetItem, QWidget, QDialog, QLabel, \
    QLineEdit, QPushButton, QVBoxLayout, QHBoxLayout, QFileDialog, QMessageBox, \
    QComboBox, QDateEdit, QTextEdit, QSizePolicy, QGridLayout

DB_DIR = "database"
# Функция за създаване на база данни и таблици
DB_PATH = "database/protocols.db"


def initialize_database():
    if not os.path.exists(DB_DIR):
        os.makedirs(DB_DIR)
    # Проверка дали файлът за базата данни съществува
    if not os.path.exists(DB_PATH):
        open(DB_PATH, 'w').close()  # Създаване на празен файл за базата
        print("Създаден е нов файл за базата данни.")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute('''CREATE TABLE IF NOT EXISTS devices (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_name TEXT,
        company_address TEXT,
        company_manager TEXT,
        device_address TEXT,
        phone_number TEXT,
        device TEXT,
        serial_number TEXT UNIQUE,
        eik TEXT,
        fdrid TEXT
    )''')

    cursor.execute('''CREATE TABLE IF NOT EXISTS protocols (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        serial_number TEXT,
        problem_description TEXT,
        date TEXT,
        FOREIGN KEY(serial_number) REFERENCES devices(serial_number)
    )''')

    conn.commit()
    conn.close()




#Create MainWindow for the start screen
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.edit_device_window = None
        self.new_device_window = NewDeviceWindow(self)  # Създаване на инстанция на прозореца
        self.edit_protocol_window = None
        self.new_protocol_window = NewProtocolWindow(self)  # Създаване на инстанция на прозореца
        self.setWindowTitle("Протоколи")
        self.setGeometry(100, 100, 1000, 700)
        self.devices_window = DevicesWindow(self)
        self.devices_table = DevicesWindow(self)  # Създаване на инстанция на прозореца за продукти
        self.protocols_window = ProtocolsWindow(self)
        self.protocols_table = ProtocolsWindow(self)

        menubar = self.menuBar()

        operations_menu = menubar.addMenu("Операции")

        new_device_action = QAction("Ново устройство", self)
        operations_menu.addAction(new_device_action)

        edit_device_action = QAction("Редактирай устройство", self)
        operations_menu.addAction(edit_device_action)

        # new_protocol_action = QAction("Нов протокол", self)
        # operations_menu.addAction(new_protocol_action)

        edit_protocol_action = QAction("Редактирай протокол", self)
        operations_menu.addAction(edit_protocol_action)

        work_with_devices_action = QAction("Работа с устройства", self)
        operations_menu.addAction(work_with_devices_action)
        work_with_devices_action.setShortcut("Ctrl+D")  # Задаване на клавишна комбинация Ctrl+P
        work_with_devices_action.setStatusTip(
            "Отваря прозореца за управление на устройствата")  # Задаване на подсказка при задържане на мишката върху действието
        work_with_devices_action.triggered.connect(self.open_devices_window)

        work_with_protocls_action = QAction("Работа с протоколи", self)
        operations_menu.addAction(work_with_protocls_action)
        work_with_protocls_action.setShortcut("Ctrl+P")  # Задаване на клавишна комбинация Ctrl+P
        work_with_protocls_action.setStatusTip(
            "Отваря прозореца за управление на протоколите")  # Задаване на подсказка при задържане на мишката върху действието
        work_with_protocls_action.triggered.connect(self.open_protocols_window)

        new_device_action.triggered.connect(self.open_new_device_window)
        edit_device_action.triggered.connect(self.edit_device)

        # new_protocol_action.triggered.connect(self.open_protocols_window)
        edit_protocol_action.triggered.connect(self.edit_protocol)
        #
        # settings_menu = menubar.addMenu("Настройки")
        # # Тук ще добавим действия към меню "Настройки"

        for_us = QAction("За програмата", self)
        for_us.triggered.connect(self.open_for_us_window)

        help_menu = menubar.addMenu("Помощ")
        help_menu.addAction(for_us)
        # Тук ще добавим действия към меню "Помощ"

        #Тук ще сложим редовете
        self.company_name_label = QLabel("Име на фирма:")
        self.company_name_input = QLineEdit()

        self.device_model_label = QLabel("Устройство модел:")
        self.device_model_input = QLineEdit()

        self.serial_number_label = QLabel("Серйен номер:")
        self.serial_number_input_1 = QLineEdit()

        self.bulstat_label = QLabel("Булстат:")
        self.bulstat_input = QLineEdit()

        self.fdrid_label = QLabel("FDRID:")
        self.fdrid_input = QLineEdit()

        self.company_address_label = QLabel("Адрес на фирмата:")
        self.company_address_input = QLineEdit()

        self.company_manager_label = QLabel("Управител:")
        self.company_manager_input = QLineEdit()

        self.device_address_label = QLabel("Адрес на устройството:")
        self.device_address_input = QLineEdit()

        # Създаване на полета за въвеждане на описание
        self.phone_label = QLabel("Телефон:")
        self.phone_input = QLineEdit()

        self.all_ids = []

        self.load_protocols()
        if len(self.all_ids) > 0:
            id_nums = len(self.all_ids)
            new_id_num = id_nums + 1
            self.id_text = str(new_id_num)
        else:
            self.id_text = "1"


        self.devices_table = QTableWidget()
        self.devices_table.setColumnCount(4)  # Задайте броя на колоните според вашите нужди
        self.devices_table.setHorizontalHeaderLabels(
            ["Име на Фирма", "Устройство модел", "Сериен номер" "Булстат", "FDRID", "Адрес на фирмата", "Управител", "Адрес на устройството", "Телефон"])  # Задаване на заглавия на колоните

        self.protocols_table = QTableWidget()
        self.protocols_table.setColumnCount(4)  # Задайте броя на колоните според вашите нужди
        self.protocols_table.setHorizontalHeaderLabels(
            ["ID", "Сериен номер", "Описание на проблема" "Дата"])  # Задаване на заглавия на колоните


        self.device_select_label = QLabel("Изберете устройство")
        self.device_select_input = QPushButton("Изберете устройство")
        self.device_select_input.clicked.connect(self.open_devices_window)
        self.device_select_input.setShortcut("F4")  # Задаване на клавишна комбинация Ctrl+P

        self.clear_button = QPushButton("Изчисти полетата")
        self.clear_button.clicked.connect(self.clear_fields)


        # Създаване на основен уиджет за прозореца
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        for widget in [self.company_name_input, self.company_name_input, self.device_model_label, self.device_model_input, self.serial_number_label,
                       self.serial_number_input_1, self.bulstat_label, self.bulstat_input, self.fdrid_label, self.fdrid_input, self.company_address_label,
                       self.company_address_input, self.company_manager_label, self.company_manager_input, self.device_address_label,  self.device_address_input,
                       self.phone_input, self.device_select_label, self.device_select_input
                       ]:
            widget.setFixedWidth(200)

        # Основен layout
        self.main_layout = QGridLayout(central_widget)

        self.main_layout.addWidget(self.company_name_label, 0, 0)
        self.main_layout.addWidget(self.company_name_input, 0, 1)
        self.main_layout.addWidget(self.device_model_label, 1, 0)
        self.main_layout.addWidget(self.device_model_input, 1, 1)
        self.main_layout.addWidget(self.serial_number_label, 2, 0)
        self.main_layout.addWidget(self.serial_number_input_1, 2, 1)
        self.main_layout.addWidget(self.bulstat_label, 3, 0)
        self.main_layout.addWidget(self.bulstat_input, 3, 1)
        self.main_layout.addWidget(self.fdrid_label, 4, 0)
        self.main_layout.addWidget(self.fdrid_input, 4, 1)
        self.main_layout.addWidget(self.company_address_label, 5, 0)
        self.main_layout.addWidget(self.company_address_input, 5, 1)
        self.main_layout.addWidget(self.company_manager_label, 6, 0)
        self.main_layout.addWidget(self.company_manager_input, 6, 1)
        self.main_layout.addWidget(self.device_address_label, 7, 0)
        self.main_layout.addWidget(self.device_address_input, 7, 1)
        self.main_layout.addWidget(self.phone_label, 8, 0)
        self.main_layout.addWidget(self.phone_input, 8, 1)
        self.main_layout.addWidget(self.clear_button, 9, 0)
        self.main_layout.addWidget(self.device_select_input, 9, 1)

        # main_layout.addWidget(self.devices_table)
        # main_layout.addWidget(self.protocols_table)
        self.setLayout(self.main_layout)

        self.load_devices()  # Зареждане на устройствата в таблицата
        self.load_protocols()

    def open_devices_window(self):
        self.devices_window.exec()

    def open_protocols_window(self):
        self.protocols_window.exec()

    def open_new_device_window(self):
        result = self.new_device_window.exec()  # Показване на прозореца и очакване за резултат
        if result == QDialog.DialogCode.Accepted:  # Проверка дали е натиснат бутонът "Запази"
            print("Устройството е запазено.")
        else:
            print("Действието е отказано.")

    def edit_device(self, device_id):
        try:
            self.edit_device_window = EditDeviceWindow(device_id, self)
            result = self.edit_device_window.exec()
            if result == QDialog.DialogCode.Accepted:
                # Обновяване на списъка с продукти
                self.load_devices()  # Това е функция, която трябва да бъде дефинирана за зареждане на продуктите от базата данни
        except Exception as e:
            print(f"Грешка при редактиране на устройство: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при редактирането на продукта: {e}")

    def edit_protocol(self, protocol_id):
        try:
            self.edit_protocol_window = EditProtocolWindow(protocol_id, self)
            result = self.edit_protocol_window.exec()
            if result == QDialog.DialogCode.Accepted:
                # Обновяване на списъка с продукти
                self.load_protocols()  # Това е функция, която трябва да бъде дефинирана за зареждане на продуктите от базата данни
        except Exception as e:
            print(f"Грешка при редактиране на протоколите: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при редактирането на протоколите: {e}")

    def load_devices(self):
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid FROM devices"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            devices = cur.fetchall()

            self.devices_table.setRowCount(0)  # Изчистване на таблицата
            self.devices_table.setColumnCount(9)  # Задаване на броя на колоните (id, name, barcode, price, image)
            self.devices_table.setHorizontalHeaderLabels(
                ["Име на Фирма", "Устройство модел", "Сериен номер" "Булстат", "FDRID", "Адрес на фирмата", "Управител", "Адрес на устройството", "Телефон"])  # Задаване на заглавките на колоните

            for device in devices:
                row_num = self.devices_table.rowCount()
                self.devices_table.insertRow(row_num)

                for i, data in enumerate(device):
                    item = QTableWidgetItem(str(data))  # Създаване на елемент от таблицата със стойността на данните
                    self.devices_table.setItem(row_num, i, item)  # Добавяне на елемента към съответната клетка
            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на устройствата: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на устройствата: {e}")

    def load_protocols(self):
        # id
        # serial_number
        # problem_description
        # date
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT id, serial_number, problem_description, date FROM protocols"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            protocols = cur.fetchall()
            self.all_protocols = []  # Изчистване на списъка с оригиналните данни

            self.protocols_table.setRowCount(0)  # Изчистване на таблицата
            # self.protocols_table.insertColumn(0)
            for protocol in protocols:
                row_num = self.protocols_table.rowCount()
                self.protocols_table.insertRow(row_num)
                # checkbox = QCheckBox()
                # self.protocols_table.setCellWidget(row_num, 0, checkbox)  # Добавяне на checkbox към първата колона
                protocol_data = {}  # Речник за съхранение на данните за продукта
                for i, data in enumerate(protocol):
                    item = QTableWidgetItem(str(data))
                    self.protocols_table.setItem(row_num, i, item)  #i+1
                    protocol_data[self.column_mapping[i]] = data  # Добавяне на данните в речника

                self.all_protocols.append(protocol_data)  # Добавяне на данните за продукта към списъка с оригиналните данни

                # Задаване на подравняване вдясно за елементите от колоните с цена и количество
                item = self.protocols_table.item(row_num, 0)  # Вземане на елемента от колоната с баркода
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране
                item = self.protocols_table.item(row_num, 2)  # Вземане на елемента от колоната с цена
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

                item = self.protocols_table.item(row_num, 3)  # Вземане на елемента от колоната с количество
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

            cur.close()
            conn.close()
            self.protocols_table.setColumnWidth(1, 200)  # Увеличава ширината на втората колона (индекс 1) на 200 пиксела

        except Exception as e:
            print(f"Грешка при зареждане на протоколите: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на протоколите: {e}")

    def load_device_data(self):
        instance_device_id = DevicesWindow()
        device_id = instance_device_id.choose_product()
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за извличане на информация за продукта
            sql = "SELECT company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid FROM devices WHERE serial_number = ?"
            cur.execute(sql, (device_id,))
            device_data = cur.fetchone()

            if device_data:
                company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid = device_data

                self.company_name_input.setText(company_name)
                self.device_model_input.setText(device_model)
                self.serial_number_input.setText(serial_number)
                self.bulstat_input.setText(bulstat)
                self.fdrid_input.setText(fdrid)
                self.company_address_input.setText(company_address)
                self.company_manager_input.setText(company_manager)
                self.device_address_input.setText(device_address)
                self.phone_input.setText(phone)

            else:
                QMessageBox.warning(self, "Грешка", "Устройството не е намерен.")
                self.close()

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на данни за устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на данните за устройството: {e}")

    def load_protocols(self):
        # id
        # serial_number
        # problem_description
        # date
        try:
            conn = sqlite3.connect(DB_PATH)# Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT id FROM protocols"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            ids = cur.fetchall()
            self.all_ids = []  # Изчистване на списъка с оригиналните данни

            # self.protocols_table.insertColumn(0)
            for id in ids:
                self.all_ids.append(id)  # Добавяне на данните за продукта към списъка с оригиналните данни
            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на ид-тата: {e}")

    def id_tex(self):
        return self.id_text

    def clear_fields(self):
        instance = DevicesWindow()
        self.company_name_input.clear()
        self.device_model_input.clear()
        self.serial_number_input_1.clear()
        self.bulstat_input.clear()
        self.fdrid_input.clear()
        self.company_address_input.clear()
        self.company_manager_input.clear()
        self.device_address_input.clear()
        self.phone_input.clear()

    def open_for_us_window(self):
        """Отваря малък прозорец с информация за автора и версията."""
        QMessageBox.information(self, "За програмата",
                                "Създадена от: 'ВладиКомпютърс 2000' ЕООД\nВерсия: 2.0\n© 2025 Всички права запазени.")


#The window for the work with the devices
class DevicesWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.column_mapping = {
            0: "company_name",
            1: "device_model",
            2: "serial_number",
            3: "bulstat",
            4: "fdrid",
            5: "company_address",
            6: "company_manager",
            7: "device_address",
            8: "phone",
        }

        self.setWindowTitle("Устройства")
        self.setGeometry(100, 100, 1000, 600)
        # self.setWindowIcoN("")

        # Създаване на таблица за продуктите
        self.devices_table = QTableWidget()
        self.devices_table.setColumnCount(9)  # Задайте броя на колоните според вашите нужди
        self.devices_table.setHorizontalHeaderLabels(
            ["Име на Фирма", "Устройство модел", "Сериен номер", "Булстат", "FDRID", "Адрес на фирмата", "Управител", "Адрес на устройството", "Телефон"])  # Задаване на заглавия на колоните
        self.load_devices()  # Зареждане на продуктите в таблицата

        # Създаване на бутон "Изтрий"
        self.choose_button = QPushButton("Избери")
        self.choose_button.clicked.connect(self.choose_product)

        # Създаване на бутони
        self.new_button = QPushButton("Нов")
        self.new_button.clicked.connect(self.open_new_devices_window)

        self.edit_button = QPushButton("Редактирай")
        self.edit_button.clicked.connect(self.open_edit_device_window)

        # Създаване на бутон "Изтрий"
        self.delete_button = QPushButton("Изтрий")
        self.delete_button.clicked.connect(self.delete_product)


        # Създаване на бутон "Обнови"
        self.refresh_button = QPushButton("Обнови")
        self.refresh_button.clicked.connect(self.load_devices)

        self.print_button = QPushButton("Отпечатване")
        self.print_button.clicked.connect(self.print_table)

        # Създаване на поле за търсене
        self.search_input = QLineEdit()
        self.search_input.textChanged.connect(self.search_devices)

        # Създаване на падащо меню за критерий за сортиране
        self.sort_criterion_combo = QComboBox()
        self.sort_criterion_combo.addItems(["Име на Фирма", "Сериен номер", "Булстат"])  # Добавяне на възможните критерии
        self.sort_criterion_combo.currentIndexChanged.connect(
            self.sort_devices)  # Свързване на сигнала currentIndexChanged с функцията за сортиране

        # Създаване на падащо меню за посока на сортиране
        self.sort_direction_combo = QComboBox()
        self.sort_direction_combo.addItems(["Възходящо", "Низходящо"])  # Добавяне на възможните посоки
        self.sort_direction_combo.currentIndexChanged.connect(
            self.sort_devices)  # Свързване на сигнала currentIndexChanged с функцията за сортиране
        #
        # Създаване на бутон "Импортиране"
        self.import_button = QPushButton("Импортиране")
        self.import_button.clicked.connect(self.import_from_csv)

        # Създаване на бутон "Изтрий избраните"
        self.delete_selected_button = QPushButton("Изтрий избраните")
        self.delete_selected_button.clicked.connect(self.delete_selected_products)

        self.serial_number_filter_input = QLineEdit()
        self.serial_number_filter_input.textChanged.connect(self.check_input)

        self.company_name_filter_input = QLineEdit()
        self.company_name_filter_input.textChanged.connect(self.check_input)

        self.bulstat_filter_input = QLineEdit()
        self.bulstat_filter_input.textChanged.connect(self.check_input)

        # Добавяне на полетата за филтриране към layout-а
        filter_layout = QVBoxLayout()
        filter_layout.addWidget(QLabel("Име на фирма:"))
        filter_layout.addWidget(self.company_name_filter_input)
        filter_layout.addWidget(QLabel("Сериен номер:"))
        filter_layout.addWidget(self.serial_number_filter_input)
        filter_layout.addWidget(QLabel("Булстат:"))
        filter_layout.addWidget(self.bulstat_filter_input)

        # Добавяне на падащите менюта към layout-а
        sort_layout = QHBoxLayout()
        sort_layout.addWidget(self.sort_criterion_combo)
        sort_layout.addWidget(self.sort_direction_combo)

        # Създаване на бутон "Експортиране"
        self.export_button = QPushButton("Експортиране")
        self.export_button.clicked.connect(self.export_to_csv)

        # Свързване на сигнала cellDoubleClicked с функцията за редактиране
        self.devices_table.cellDoubleClicked.connect(self.choose_product)

        # Разположение на елементите
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.new_button)
        button_layout.addWidget(self.edit_button)
        button_layout.addWidget(self.choose_button)
        # Добавяне на бутона "Изтрий" към layout-а
        button_layout.addWidget(self.delete_button)
        button_layout.addWidget(self.refresh_button)
        # Добавяне на бутона "Отпечатване" към layout-а
        button_layout.addWidget(self.print_button)
        button_layout.addWidget(self.export_button)
        # Добавяне на бутона "Импортиране" към layout-а
        button_layout.addWidget(self.import_button)
        # Добавяне на бутона "Изтрий избраните" към layout-а
        # button_layout.addWidget(self.delete_selected_button)

        main_layout = QVBoxLayout()
        # main_layout.addWidget(self.search_input)
        main_layout.addWidget(self.devices_table)
        main_layout.addLayout(button_layout)
        main_layout.addLayout(sort_layout)
        main_layout.addLayout(filter_layout)

        self.setLayout(main_layout)

        self.all_devices = []  # Създаване на празен списък за оригиналните данни

    def choose_product(self, row, column):
        selected_rows = self.devices_table.selectedItems()
        if not selected_rows:
            QMessageBox.warning(self, "Грешка", "Моля, изберете устройство.")
            return

        column_count = self.devices_table.columnCount()
        row_data = []

        for col in range(column_count):
            item = self.devices_table.item(row, col)  # Вземаме клетката
            row_data.append(item.text() if item is not None else "")  # Проверяваме дали е None

        print(f"Избран ред ({row}): {row_data}")  # Дебъгване

        # Ако имаш нужда от ID (първа колона)
        device_id = row_data[2] if row_data else ""

        print(f"ID : {device_id}")

        # device_id = selected_rows[2].text()
        self.device_id = device_id
        self.load_device_data()

    def load_device_data(self):
        instance_main = self.parent()
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за извличане на информация за продукта
            sql = "SELECT company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid FROM devices WHERE serial_number = ?"
            cur.execute(sql, (self.device_id,))
            self.device_data = cur.fetchone()

            if self.device_data:
                company_name, company_address, company_manager, device_address, phone_number, device, self.serial_number, eik, fdrid = self.device_data

                instance_main.company_name_input.setText(company_name)
                instance_main.device_model_input.setText(device)
                instance_main.serial_number_input_1.setText(self.serial_number)
                instance_main.bulstat_input.setText(str(eik))
                instance_main.fdrid_input.setText(str(fdrid))
                instance_main.company_address_input.setText(company_address)
                instance_main.company_manager_input.setText(company_manager)
                instance_main.device_address_input.setText(device_address)
                instance_main.phone_input.setText(str(phone_number))

            else:
                QMessageBox.warning(self, "Грешка", "Устройството не е намерен.")
                self.close()

            cur.close()
            conn.close()
            self.result = instance_main.id_text

            # Създаване на полета за въвеждане на данни
            instance_main.id_label = QLabel("ID:")
            instance_main.id_input = QLineEdit()
            instance_main.id_input.setText(self.result)  # Задаване на текст
            instance_main.id_input.setReadOnly(True)

            instance_main.serial_number_label = QLabel("Серйен номер:")
            instance_main.serial_number_input = QLineEdit()
            instance_main.serial_number_input.setText(self.serial_number)  # Задаване на текст
            instance_main.serial_number_input.setReadOnly(True)

            instance_main.problem_description_label = QLabel("Описание на проблема:")
            instance_main.problem_description_input = QTextEdit()
            instance_main.problem_description_input.setMinimumSize(200, 80)  # Минимален размер
            instance_main.problem_description_input.setMaximumSize(400, 150)  # Максимален размер

            # Създаване на полета за въвеждане на описание
            instance_main.date_label = QLabel("Дата:")
            instance_main.date_input = QDateEdit()
            instance_main.date_input.setCalendarPopup(True)  # Позволява изскачащ календар
            instance_main.date_input.setDate(QDate.currentDate())  # Задава днешната дата по подразбиране

            instance_main.save_button_input = QPushButton("Запази протокола")
            instance_main.save_button_input.clicked.connect(self.save_protocol)

            instance_main.generate_input = QPushButton("Запази протокола")
            instance_main.generate_input.clicked.connect(self.generate_protocol)

            instance_main.main_layout.addWidget(instance_main.id_label, 0, 2)
            instance_main.main_layout.addWidget(instance_main.id_input, 0, 3)
            instance_main.main_layout.addWidget(instance_main.serial_number_label, 1, 2)
            instance_main.main_layout.addWidget(instance_main.serial_number_input, 1, 3)
            instance_main.main_layout.addWidget(instance_main.problem_description_label, 2, 2)
            instance_main.main_layout.addWidget(instance_main.problem_description_input, 2, 3)
            instance_main.main_layout.addWidget(instance_main.date_label, 3, 2)
            instance_main.main_layout.addWidget(instance_main.date_input, 3, 3)

            instance_main.main_layout.addWidget(instance_main.save_button_input, 4, 2)
            instance_main.main_layout.addWidget(instance_main.generate_input, 4, 3)

            print(instance_main.problem_description_input.toPlainText())
            self.date = instance_main.date_input.date().toString("yyyy-MM-dd")  # Примерен формат

            self.close()

        except Exception as e:
            print(f"Грешка при зареждане на данни за устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на данните за устройството: {e}")

    def check_input(self):
        if not self.company_name_filter_input.text().strip() and not self.serial_number_filter_input.text().strip() and not self.bulstat_filter_input.text().strip():
            self.load_devices()
        else:
            self.load_devices()
            self.filter_products()

    def load_devices(self):
        #"Име на Фирма", "Устройство модел", "Сериен номер" "Булстат", "FDRID", "Адрес на фирмата", "Управител", "Адрес на устройството", "Телефон"
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT company_name, device, serial_number, eik, fdrid, company_address, company_manager, device_address, phone_number FROM devices"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            devices = cur.fetchall()
            self.all_devices = []  # Изчистване на списъка с оригиналните данни

            self.devices_table.setRowCount(0)  # Изчистване на таблицата
            # self.protocols_table.insertColumn(0)
            for device in devices:
                row_num = self.devices_table.rowCount()
                self.devices_table.insertRow(row_num)
                # checkbox = QCheckBox()
                # self.protocols_table.setCellWidget(row_num, 0, checkbox)  # Добавяне на checkbox към първата колона
                device_data = {}  # Речник за съхранение на данните за продукта
                for i, data in enumerate(device):
                    item = QTableWidgetItem(str(data))
                    self.devices_table.setItem(row_num, i, item)  #i+1
                    device_data[self.column_mapping[i]] = data  # Добавяне на данните в речника

                self.all_devices.append(device_data)  # Добавяне на данните за продукта към списъка с оригиналните данни

                # Задаване на подравняване вдясно за елементите от колоните с цена и количество
                item = self.devices_table.item(row_num, 0)  # Вземане на елемента от колоната с баркода
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране
                item = self.devices_table.item(row_num, 2)  # Вземане на елемента от колоната с цена
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

                item = self.devices_table.item(row_num, 3)  # Вземане на елемента от колоната с количество
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

            cur.close()
            conn.close()
            self.devices_table.setColumnWidth(1, 200)  # Увеличава ширината на втората колона (индекс 1) на 200 пиксела

        except Exception as e:
            print(f"Грешка при зареждане на устройствата: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на устройствата: {e}")

    def open_new_devices_window(self):
        self.new_devices_window = NewDeviceWindow(self)
        result = self.new_devices_window.exec()
        if result == QDialog.DialogCode.Accepted:
            self.load_devices()

    def open_edit_device_window(self):
        selected_rows = self.devices_table.selectedItems()
        if not selected_rows:
            QMessageBox.warning(self, "Грешка", "Моля, изберете устройство за редактиране.")
            return

        device_id = selected_rows[2].text()  # Вземане на ID на продукта от избрания ред

        self.edit_devices_window = EditDeviceWindow(device_id, self)
        result = self.edit_devices_window.exec()
        if result == QDialog.DialogCode.Accepted:
            self.load_devices()

    def search_devices(self, text):
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT company_name, device, serial_number, eik, fdrid, company_address, company_manager, device_address, phone_number FROM devices WHERE serial_number LIKE ?"  # Търсене по име (ILIKE е case-insensitive)
            cur.execute(sql, ('%' + text + '%',))  # Добавяне на wildcard символи (%) за търсене на подstring
            devices = cur.fetchall()

            self.devices_table.setRowCount(0)  # Изчистване на таблицата
            for device in devices:
                row_num = self.devices_table.rowCount()
                self.devices_table.insertRow(row_num)
                for i, data in enumerate(device):
                    item = QTableWidgetItem(str(data))
                    self.devices_table.setItem(row_num, i, item)

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при търсене на устройствата: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при търсенето на устройствата: {e}")

    def sort_devices(self):
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Вземане на избрания критерий и посока на сортиране
            criterion = self.sort_criterion_combo.currentText()
            direction = self.sort_direction_combo.currentText()

            # Изпълнение на SQL заявка за сортиране на продуктите
            if criterion == "Име на Фирма":
                sql = "SELECT company_name, device, serial_number, eik, fdrid, company_address, company_manager, device_address, phone_number FROM devices ORDER BY company_name "  # SQL заявка за извличане на данните за продуктите
            elif criterion == "Сериен номер":
                sql = "SELECT company_name, device, serial_number, eik, fdrid, company_address, company_manager, device_address, phone_number FROM devices ORDER BY serial_number "  # SQL заявка за извличане на данните за продуктите
            elif criterion == "Булстат":
                sql = "SELECT company_name, device, serial_number, eik, fdrid, company_address, company_manager, device_address, phone_number FROM devices ORDER BY eik "  # SQL заявка за извличане на данните за продуктите

            if direction == "Низходящо":
                sql += "DESC"

            cur.execute(sql)
            devices = cur.fetchall()

            self.devices_table.setRowCount(0)  # Изчистване на таблицата
            for device in devices:
                row_num = self.devices_table.rowCount()
                self.devices_table.insertRow(row_num)
                for i, data in enumerate(device):
                    item = QTableWidgetItem(str(data))
                    self.devices_table.setItem(row_num, i, item)

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при сортиране на устройства: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при сортирането на устройства: {e}")

    def filter_products(self):
        try:
            # Вземане на текста от полетата за филтриране
            company_name_filter = self.company_name_filter_input.text().lower()
            serial_number_filter = self.serial_number_filter_input.text().lower()
            bulstat_filter = self.bulstat_filter_input.text()

            if company_name_filter == '':
                company_name_filter = "none"
            if serial_number_filter == '':
                serial_number_filter = "none"
            if bulstat_filter == '':
                bulstat_filter = "none"

            print(self.all_devices)
            # Филтриране на данните
            filtered_devices = []
            for device in self.all_devices:
                if company_name_filter in device["company_name"].lower():
                    filtered_devices.append(device)
                elif serial_number_filter in device["serial_number"].lower():
                    filtered_devices.append(device)
                elif bulstat_filter in str(device["bulstat"]):
                    filtered_devices.append(device)
                # # Използваме копие на оригиналните данни
                # if (company_name_filter in device["company_name"].lower() or
                #         serial_number_filter in device["serial_number"].lower() or
                #         (bulstat_filter in str(device["eik"]))):
                #     filtered_devices.append(device)

            # Обновяване на таблицата
            self.devices_table.setRowCount(0)  # Изчистване на таблицата
            for device in filtered_devices:
                row_num = self.devices_table.rowCount()
                self.devices_table.insertRow(row_num)
                for i, (col_name, value) in enumerate(device.items()):
                    item = QTableWidgetItem(str(value))
                    self.devices_table.setItem(row_num, i, item)

        except Exception as e:
            print(f"Грешка при филтриране на устройства: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при филтрирането на устройствата: {e}")

    def print_table(self):
        try:
            # Създаване на печатащ документ
            printer = QPrinter()

            # Създаване на диалогов прозорец за настройка на печат
            print_dialog = QPrintDialog(printer, self)
            if print_dialog.exec() == QPrintDialog.DialogCode.Accepted:
                # Създаване на painter за рисуване върху печатащия документ
                painter = QPainter(printer)

                # Рисуване на данните от таблицата
                self.devices_table.render(painter)

                painter.end()

        except Exception as e:
            print(f"Грешка при отпечатване на таблицата: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при отпечатването на таблицата: {e}")

    # def validate_data(self, row_data):
    #     try:
    #         # Проверка за празни полета
    #         if not row_data["name"]:
    #             raise ValueError("Полето 'име' е задължително.")
    #         if not row_data["barcode"]:
    #             raise ValueError("Полето 'баркод' е задължително.")
    #         if not row_data["price"]:
    #             raise ValueError("Полето 'цена' е задължително.")
    #         if not row_data["quantity"]:
    #             raise ValueError("Полето 'количество' е задължително.")
    #
    #         # Проверка за валидни типове данни
    #         float(row_data["price"])  # Проверка дали цената е число
    #         int(row_data["quantity"])  # Проверка дали количеството е цяло число
    #
    #         return True  # Данните са валидни
    #
    #     except ValueError as e:
    #         QMessageBox.warning(self, "Грешка при валидиране", str(e))
    #         return False  # Данните не са валидни

    def import_from_csv(self):
        try:
            # Отваряне на диалогов прозорец за избор на файл
            file_path, _ = QFileDialog.getOpenFileName(self, "Избор на CSV файл", "", "CSV Files (*.csv)")

            if file_path:
                # Четене на данните от CSV файла
                with open(file_path, "r", encoding="utf-8-sig") as csvfile:
                    reader = csv.DictReader(csvfile)  # Използване на DictReader за четене на данни по колони
                    data = list(reader)

                # Добавяне на данните към таблицата
                self.devices_table.setRowCount(0)  # Изчистване на таблицата преди добавяне на нови данни
                for row_data in data:
                    row_num = self.devices_table.rowCount()
                    self.devices_table.insertRow(row_num)
                    for col_name, value in row_data.items():
                        col_index = list(self.column_mapping.keys())[list(self.column_mapping.values()).index(
                            col_name)]  # Намиране на индекса на колоната по име
                        item = QTableWidgetItem(value)
                        self.devices_table.setItem(row_num, col_index, item)

                QMessageBox.information(self, "Успех", "Данните са успешно импортирани от CSV файл.")

                # Питаме потребителя дали иска да запази данните в базата данни
                reply = QMessageBox.question(self, "Запазване в база данни",
                                             "Желаете ли да запазите импортираните данни в базата данни?",
                                             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if reply == QMessageBox.StandardButton.Yes:
                    try:
                        conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
                        cur = conn.cursor()

                        # Проверяваме дали записите вече съществуват и добавяме само новите
                        for row_data in data:
                            #     if not self.validate_data(row_data):
                            #         continue  # Преминаваме към следващия запис, ако данните не са валидни
                            # Изграждаме SQL заявка за проверка дали записът вече съществува
                            check_sql = "SELECT id FROM products WHERE barcode = ?"  # Проверяваме по баркод, можете да промените критерия
                            cur.execute(check_sql, (row_data["barcode"],))
                            existing_product = cur.fetchone()

                            if existing_product is None:  # Ако записът не съществува, го добавяме
                                # Изграждаме SQL заявка за добавяне на нов запис
                                insert_sql = "INSERT INTO products (name, barcode, price, quantity) VALUES (?, ?, ?, ?)"  # Адаптирайте колоните според вашите нужди
                                values = (row_data["name"], row_data["barcode"], row_data["price"],
                                          row_data["quantity"])  # Адаптирайте данните според вашите нужди
                                cur.execute(insert_sql, values)

                        conn.commit()
                        cur.close()
                        conn.close()

                        QMessageBox.information(self, "Успех", "Данните са успешно запазени в базата данни.")
                        self.load_devices()  # Обновяваме таблицата с продуктите след импортирането и запазването
                    except Exception as e:
                        print(f"Грешка при запазване в базата данни: {e}")
                        conn.rollback()  # Откатваме транзакцията в случай на грешка
                        QMessageBox.critical(self, "Грешка",
                                             f"Възникна грешка при запазването на данните в базата данни: {e}")


        except Exception as e:
            print(f"Грешка при импортиране от CSV файл: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при импортирането на данните от CSV файл: {e}")

    def export_to_csv(self):
        try:
            # Получаване на данните от таблицата, включително заглавките на колоните
            headers = [self.devices_table.horizontalHeaderItem(i).text() for i in
                       range(self.devices_table.columnCount())]
            data = []
            for row in range(self.devices_table.rowCount()):
                row_data = {}  # Променяме на речник, за да запазим реда на колоните
                for col in range(self.devices_table.columnCount()):
                    item = self.devices_table.item(row, col)
                    if item is not None:
                        column_name = self.column_mapping[col]  # Вземаме името на колоната от речника
                        row_data[column_name] = item.text()
                    else:
                        column_name = self.column_mapping[col]
                        row_data[column_name] = ""  # Добавяне на празен низ, ако клетката е празна
                data.append(row_data)
                print(data)

            # Отваряне на диалогов прозорец за запазване на файл
            file_path, _ = QFileDialog.getSaveFileName(self, "Запазване на CSV файл", "", "CSV Files (*.csv)")

            if file_path:
                # Запис на данните в CSV файла с UTF-8 кодиране и BOM
                with open(file_path, "w", newline="", encoding="utf-8-sig") as csvfile:
                    fieldnames = list(self.column_mapping.values())  # Вземаме имената на колоните от речника
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=",", quoting=csv.QUOTE_NONNUMERIC)
                    writer.writeheader()  # Запис на заглавките на колоните
                    writer.writerows(data)  # Запис на данните

                QMessageBox.information(self, "Успех", "Данните са успешно експортирани в CSV файл.")
                os_name = platform.system()
                # Питаме потребителя дали иска да отвори файла
                reply = QMessageBox.question(self, "Отваряне на файл", "Искате ли да отворите файла?",
                                             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if reply == QMessageBox.StandardButton.Yes:
                    # Отваряне на файла в зависимост от операционната система
                    if os_name == "Windows":
                        os.startfile(file_path)
                    elif os_name == "Darwin":  # macOS
                        subprocess.Popen(['open', file_path])
                    elif os_name == "Linux":
                        subprocess.Popen(['xdg-open', file_path])
                    else:
                        QMessageBox.warning(self, "Грешка", "Неподдържана операционна система.")

        except Exception as e:
            print(f"Грешка при експортиране в CSV файл: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при експортирането на данните в CSV файл: {e}")

    def edit_selected_product(self, row, column):
        try:
            # Вземане на ID на избрания продукт
            device_id = int(self.devices_table.item(row, 0).text())

            # Отваряне на прозореца за редактиране на продукт
            self.edit_devices_window = EditDeviceWindow(device_id, self)
            result = self.edit_devices_window.exec()

            # Обновяване на таблицата с продуктите, ако потребителят е запазил промените
            if result == QDialog.DialogCode.Accepted:
                self.load_devices()

        except Exception as e:
            print(f"Грешка при отваряне на прозореца за редактиране: {e}")
            QMessageBox.critical(self, "Грешка",
                                 f"Възникна грешка при отварянето на прозореца за редактиране: {e}")

    def delete_product(self):
        selected_items = self.devices_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Грешка", "Моля, изберете устройство за изтриване.")
            return

        device_id = selected_items[2].text()

        reply = QMessageBox.question(self, "Потвърждение", "Сигурни ли сте, че искате да изтриете това устройство?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            try:
                conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
                cur = conn.cursor()

                sql = "DELETE FROM devices WHERE serial_number = ?"
                cur.execute(sql, (device_id,))

                conn.commit()
                cur.close()
                conn.close()

                self.load_devices()  # Обновяване на таблицата с продуктите

                QMessageBox.information(self, "Успех", "Устройството е успешно изтрито.")

            except Exception as e:
                print(f"Грешка при изтриване на устройство: {e}")
                QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на устройството: {e}")

    def delete_selected_products(self):
        try:
            # Вземане на избраните продукти
            selected_products = []
            for row in range(self.devices_table.rowCount()):
                checkbox = self.devices_table.cellWidget(row, 0)
                if checkbox.isChecked():
                    product_id = int(self.devices_table.item(row, 1).text())  # ID-то е във втората колона (индекс 1)
                    selected_products.append(product_id)

            if not selected_products:
                QMessageBox.warning(self, "Грешка", "Моля, изберете продукти за изтриване.")
                return

            # Показване на диалогов прозорец за потвърждение на изтриването
            reply = QMessageBox.question(self, "Потвърждение",
                                         "Сигурни ли сте, че искате да изтриете избраните продукти?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

            if reply == QMessageBox.StandardButton.Yes:
                try:
                    conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
                    cur = conn.cursor()

                    # Изтриване на избраните продукти от базата данни
                    for product_id in selected_products:
                        sql = "DELETE FROM products WHERE id = ?"
                        cur.execute(sql, (product_id,))

                    conn.commit()
                    cur.close()
                    conn.close()

                    self.load_devices()  # Обновяване на таблицата с продуктите

                    QMessageBox.information(self, "Успех", "Избраните продукти са успешно изтрити.")

                except Exception as e:
                    print(f"Грешка при изтриване на продукти: {e}")
                    QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на продуктите: {e}")

        except Exception as e:
            print(f"Грешка при изтриване на продукти: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на продуктите: {e}")

    def save_protocol(self):
        instance_main = self.parent()
        self.problems = []

        # Тук ще добавим код за запазване на продукта в базата данни
        # Тук ще добавим код за запазване на продукта в базата данни
        id = self.result
        serial_number = self.serial_number

        self.problem_description = instance_main.problem_description_input.toPlainText()
        date = self.date

        self.problems = []
        self.problems.append(id)
        self.problems.append(self.problem_description)
        self.problems.append(date)


        if not id or not serial_number  or not date:
            QMessageBox.warning(self, "Грешка", "Моля, попълнете всички полета.")
            return

        try:
            # Свързване с базата данни
            conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за вмъкване на нов продукт
            sql = "INSERT INTO protocols (id, serial_number, problem_description, date) VALUES (?, ?, ?, ?)"
            values = (id, serial_number, self.problem_description, date)
            cur.execute(sql, values)

            # Запис на промените в базата данни
            conn.commit()

            # Затваряне на връзката с базата данни
            cur.close()
            conn.close()

            QMessageBox.information(self, "Успех", "Протоколът е успешно добавено.")
            instance_main.clear_fields()
            instance_main.serial_number_input.clear()
            instance_main.clear_fields()
            self.accept()  # Затваря прозореца след успешно запазване

        except Exception as e:
            print(f"Грешка при запазване на устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при запазването на протоколът: {e}")

    def generate_protocol(self):
        instance_main = MainWindow()

        if self.device_data:
            company_name, company_address, company_manager, device_address, phone_number, device, self.serial_number, eik, fdrid = self.device_data

        document = Document("template.docx")

        protocol_id = self.problems[0]
        date_string = self.problems[2] # Форматиране на датата
        # company_name = self.device_data[0]
        # company_address = self.device_info[1]
        # company_manager = self.device_info[2]
        # device_address = self.device_info[3]
        # phone_number = self.device_info[4]
        # device = self.device_info[5]
        # serial_number = self.device_data[6]
        problem_description = self.problems[1]
        first_row = protocol_id + '/' + date_string

        for paragraph in document.paragraphs:
            print(paragraph.text)
            if "[номер на протокол от базата данни]/[дата[дд-мм-гг]]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[номер на протокол от базата данни]/[дата[дд-мм-гг]]", first_row)
            if "[Име на фирма]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[Име на фирма]", company_name)
            if "[адрес на фирма]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[адрес на фирма]", company_address)
            if "[управител]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[управител]", company_manager)
            if "[адрес на устройството]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[адрес на устройството]", device_address)
            if "[телефонен номер]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[телефонен номер]", phone_number)
            if "[какво е оставено и име и модел]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[какво е оставено и име и модел]", device)
            if "[сериен номер]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[сериен номер]", self.serial_number)
            if "[описание на порблема]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[описание на порблема]", problem_description)
            # elif "[Сериен номер]" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("[Сериен номер]", serial_number)
            # elif "[Модел]" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("[Модел]", device[1])
            # elif "[Производител]" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("[Производител]", device[2])
            # elif "[Описание на проблема]" in paragraph.text:
            #     paragraph.text = paragraph.text.replace("[Описание на проблема]", problem_description)
        pre_file_path = f"protocol_{protocol_id}_{self.serial_number}_{date_string}.docx"

        file_path, _ = QFileDialog.getSaveFileName(self, "Запазване на Docx файл", os.path.join(os.path.expanduser("~"), "Documents", pre_file_path), "Docx Files (*.docx);; All Files (*)")
        file_path = file_path+''+pre_file_path
        document.save(f"{file_path}")

        QMessageBox.information(self, "Успех", "Протоколът беше генериран успешно.")
        self.remove_fields_auto()
        os_name = platform.system()
        # Питаме потребителя дали иска да отвори файла
        reply = QMessageBox.question(self, "Отваряне на файл", "Искате ли да отворите файла?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            # Отваряне на файла в зависимост от операционната система
            if os_name == "Windows":
                os.startfile(file_path)
            elif os_name == "Darwin":  # macOS
                subprocess.Popen(['open', file_path])
            elif os_name == "Linux":
                subprocess.Popen(['xdg-open', file_path])
            else:
                QMessageBox.warning(self, "Грешка", "Неподдържана операционна система.")

    def remove_field(self, widget):
        instance_main = self.parent()

        widget.hide()  # Скрива полето
        instance_main.main_layout.removeWidget(widget)  # Премахва го от layout-а
        widget.deleteLater()  # Изтрива го от паметта

    def remove_fields_auto(self):
        instance_main = self.parent()
        self.remove_field(instance_main.id_input)
        self.remove_field(instance_main.serial_number_input)
        self.remove_field(instance_main.problem_description_input)
        self.remove_field(instance_main.date_input)
        self.remove_field(instance_main.id_label)
        self.remove_field(instance_main.serial_number_label)
        self.remove_field(instance_main.problem_description_label)
        self.remove_field(instance_main.date_label)
        self.remove_field(instance_main.save_button_input)
        self.remove_field(instance_main.generate_input)


#window for new device
class NewDeviceWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        #["Име на Фирма", "Устройство модел", "Сериен номер" "Булстат", "FDRID", "Адрес на фирмата", "Управител", "Адрес на устройството", "Телефон"])  # Задаване на заглавия на колоните

        self.setWindowTitle("Ново устройство")

        # Създаване на полета за въвеждане на данни
        self.company_name_label = QLabel("Име на фирма:")
        self.company_name_input = QLineEdit()

        self.device_model_label = QLabel("Устройство модел:")
        self.device_model_input = QLineEdit()

        self.serial_number_label = QLabel("Серйен номер:")
        self.serial_number_input = QLineEdit()

        self.bulstat_label = QLabel("Булстат:")
        self.bulstat_input = QLineEdit()

        self.fdrid_label = QLabel("FDRID:")
        self.fdrid_input = QLineEdit()

        self.company_address_label = QLabel("Aдрес на фирмата:")
        self.company_address_input = QLineEdit()

        self.company_manager_label = QLabel("Управител:")
        self.company_manager_input = QLineEdit()

        self.device_address_label = QLabel("Адрес на устройството:")
        self.device_address_input = QLineEdit()

        # Създаване на полета за въвеждане на описание
        self.phone_label = QLabel("Телефон:")
        self.phone_input = QLineEdit()


        # Създаване на бутон за запазване
        self.save_button = QPushButton("Запази")
        self.save_button.clicked.connect(self.save_device)

        # Създаване на бутон за отказ
        self.cancel_button = QPushButton("Отказ")
        self.cancel_button.clicked.connect(self.close)

        # Разположение на елементите в прозореца
        layout = QVBoxLayout()
        layout.addWidget(self.company_name_label)
        layout.addWidget(self.company_name_input)
        layout.addWidget(self.device_model_label)
        layout.addWidget(self.device_model_input)
        layout.addWidget(self.serial_number_label)
        layout.addWidget(self.serial_number_input)
        layout.addWidget(self.bulstat_label)
        layout.addWidget(self.bulstat_input)
        layout.addWidget(self.fdrid_label)
        layout.addWidget(self.fdrid_input)
        layout.addWidget(self.company_address_label)
        layout.addWidget(self.company_address_input)
        layout.addWidget(self.company_manager_label)
        layout.addWidget(self.company_manager_input)
        layout.addWidget(self.device_address_label)
        layout.addWidget(self.device_address_input)
        layout.addWidget(self.phone_label)
        layout.addWidget(self.phone_input)
        layout.addWidget(self.save_button)
        layout.addWidget(self.cancel_button)
        self.setLayout(layout)

    # def upload_image(self):
    #     try:
    #         # Отваряне на диалогов прозорец за избиране на файл
    #         # options = QFileDialog.Options()
    #         file_path, _ = QFileDialog.getOpenFileName(self, "Изберете снимка", "", "Image Files (*.png *.jpg *.jpeg)")
    #
    #         if file_path:
    #             # Зареждане на снимката и показване в QLabel
    #             pixmap = QPixmap(file_path)
    #             pixmap = pixmap.scaled(100, 100)  # Мащабиране на снимката до размера на QLabel
    #             self.image_label.setPixmap(pixmap)
    #             self.image_path = file_path  # Запазване на пътя към снимката
    #     except Exception as e:
    #         print(f"Грешка при качване на снимка: {e}")
    #         QMessageBox.critical(self, "Грешка", f"Възникна грешка при качването на снимката: {e}")

    def save_device(self):
        # Тук ще добавим код за запазване на продукта в базата данни
        company_name = self.company_name_input.text()
        device_model = self.device_model_input.text()
        serial_number = self.serial_number_input.text()
        bulstat = self.bulstat_input.text()
        fdrid = self.fdrid_input.text()
        company_address = self.company_address_input.text()
        company_manager = self.company_manager_input.text()
        device_address = self.device_address_input.text()
        phone = self.phone_input.text()


        if not company_name or not device_model or not serial_number or not bulstat or not fdrid or not company_address or not company_manager or not device_address or not phone:
            QMessageBox.warning(self, "Грешка", "Моля, попълнете всички полета.")
            return

        try:
            # Свързване с базата данни
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за вмъкване на нов продукт
            sql = "INSERT INTO devices (company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)"
            values = (company_name, company_address, company_manager, device_address, phone, device_model, serial_number, bulstat, fdrid)
            cur.execute(sql, values)

            # Запис на промените в базата данни
            conn.commit()

            # Затваряне на връзката с базата данни
            cur.close()
            conn.close()

            QMessageBox.information(self, "Успех", "Устройството е успешно добавено.")
            self.accept()  # Затваря прозореца след успешно запазване

        except Exception as e:
            print(f"Грешка при запазване на устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при запазването на устройството: {e}")


#window for edit device
class EditDeviceWindow(QDialog):
    def __init__(self, device_id, parent=None):
        super().__init__(parent)

        self.setWindowTitle("Редактиране на устройство")
        self.device_id = device_id  # Запазване на ID на продукта

        # Създаване на полета за въвеждане на данни
        self.company_name_label = QLabel("Име на фирма:")
        self.company_name_input = QLineEdit()

        self.device_model_label = QLabel("Устройство модел:")
        self.device_model_input = QLineEdit()

        self.serial_number_label = QLabel("Серйен номер:")
        self.serial_number_input = QLineEdit()

        self.bulstat_label = QLabel("Булстат:")
        self.bulstat_input = QLineEdit()

        self.fdrid_label = QLabel("FDRID:")
        self.fdrid_input = QLineEdit()

        self.company_address_label = QLabel("Адрес на фирмата:")
        self.company_address_input = QLineEdit()

        self.company_manager_label = QLabel("Управител:")
        self.company_manager_input = QLineEdit()

        self.device_address_label = QLabel("Адрес на устройството:")
        self.device_address_input = QLineEdit()

        # Създаване на полета за въвеждане на описание
        self.phone_label = QLabel("Телефон:")
        self.phone_input = QLineEdit()

        # Създаване на бутон за запазване
        self.save_button = QPushButton("Запази")
        self.save_button.clicked.connect(self.save_device)

        # Създаване на бутон за отказ
        self.cancel_button = QPushButton("Отказ")
        self.cancel_button.clicked.connect(self.close)

        # Разположение на елементите в прозореца
        layout = QVBoxLayout()
        layout.addWidget(self.company_name_label)
        layout.addWidget(self.company_name_input)
        layout.addWidget(self.device_model_label)
        layout.addWidget(self.device_model_input)
        layout.addWidget(self.serial_number_label)
        layout.addWidget(self.serial_number_input)
        layout.addWidget(self.bulstat_label)
        layout.addWidget(self.bulstat_input)
        layout.addWidget(self.fdrid_label)
        layout.addWidget(self.fdrid_input)
        layout.addWidget(self.company_address_label)
        layout.addWidget(self.company_address_input)
        layout.addWidget(self.company_manager_label)
        layout.addWidget(self.company_manager_input)
        layout.addWidget(self.device_address_label)
        layout.addWidget(self.device_address_input)
        layout.addWidget(self.phone_label)
        layout.addWidget(self.phone_input)
        layout.addWidget(self.save_button)
        layout.addWidget(self.cancel_button)
        self.setLayout(layout)

        # Зареждане на информацията за продукта от базата данни
        self.load_device_data()

    def load_device_data(self):
        try:
            conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за извличане на информация за продукта
            sql = "SELECT company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid FROM devices WHERE serial_number = ?"
            cur.execute(sql, (self.device_id,))
            device_data = cur.fetchone()

            if device_data:
                company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid = device_data

                self.company_name_input.setText(company_name)
                self.device_model_input.setText(device)
                self.serial_number_input.setText(serial_number)
                self.bulstat_input.setText(str(eik))
                self.fdrid_input.setText(str(fdrid))
                self.company_address_input.setText(company_address)
                self.company_manager_input.setText(company_manager)
                self.device_address_input.setText(device_address)
                self.phone_input.setText(str(phone_number))

            else:
                QMessageBox.warning(self, "Грешка", "Устройството не е намерен.")
                self.close()

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на данни за устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на данните за устройството: {e}")

    # def upload_image(self):
    #     try:
    #         # Отваряне на диалогов прозорец за избиране на файл
    #         # options = QFileDialog.Options()
    #         file_path, _ = QFileDialog.getOpenFileName(self, "Изберете снимка", "", "Image Files (*.png *.jpg *.jpeg)")
    #
    #         if file_path:
    #             # Зареждане на снимката и показване в QLabel
    #             pixmap = QPixmap(file_path)
    #             pixmap = pixmap.scaled(100, 100)  # Мащабиране на снимката до размера на QLabel
    #             self.image_label.setPixmap(pixmap)
    #             self.image_path = file_path  # Запазване на пътя към снимката
    #     except Exception as e:
    #         print(f"Грешка при качване на снимка: {e}")
    #         QMessageBox.critical(self, "Грешка", f"Възникна грешка при качването на снимката: {e}")

    def save_device(self):
        # Тук ще добавим код за запазване на продукта в базата данни
        company_name = self.company_name_input.text()
        device_model = self.device_model_input.text()
        serial_number = self.serial_number_input.text()
        bulstat = self.bulstat_input.text()
        fdrid = self.fdrid_input.text()
        company_address = self.company_address_input.text()
        company_manager = self.company_manager_input.text()
        device_address = self.device_address_input.text()
        phone = self.phone_input.text()

        try:
            # Свързване с базата данни
            conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за актуализиране на информацията за продукта
            # sql = "INSERT INTO devices (company_name, company_address, company_manager, device_address, phone_number, device, serial_number, eik, fdrid) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)"

            sql = "UPDATE devices SET company_name = ?, company_address = ?, company_manager = ?, device_address = ?, phone_number = ?, device = ?, serial_number = ?, eik = ?, fdrid = ? WHERE serial_number = ?"
            values = (company_name, company_address, company_manager, device_address, phone, device_model, serial_number, bulstat, fdrid, self.device_id)
            cur.execute(sql, values)

            conn.commit()
            cur.close()
            conn.close()

            QMessageBox.information(self, "Успех", "Устройството е успешно редактирано.")
            self.accept()

        except Exception as e:
            print(f"Грешка при редактиране на устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при редактирането на устройството: {e}")


#The window for the work with the protocols
class ProtocolsWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # id
        # serial_number
        # problem_description
        # date

        self.column_mapping = {
            0: "id",
            1: "serial_number",
            2: "problem_description",
            3: "date",
        }

        self.setWindowTitle("Протоколи")
        self.setGeometry(100, 100, 1000, 600)
        # self.setWindowIcoN("")

        # Създаване на таблица за продуктите
        self.protocols_table = QTableWidget()
        self.protocols_table.setColumnCount(4)  # Задайте броя на колоните според вашите нужди
        self.protocols_table.setHorizontalHeaderLabels(
            ["ID", "Сериен номер", "Описание на проблема", "Дата"])  # Задаване на заглавия на колоните
        self.load_protocols()  # Зареждане на продуктите в таблицата

        # # Създаване на бутони
        # self.new_button = QPushButton("Нов")
        # self.new_button.clicked.connect(self.open_new_protocols_window)

        self.edit_button = QPushButton("Редактирай")
        self.edit_button.clicked.connect(self.open_edit_protocols_window)

        # Създаване на бутон "Изтрий"
        self.delete_button = QPushButton("Изтрий")
        self.delete_button.clicked.connect(self.delete_product)

        # Създаване на бутон "Обнови"
        self.refresh_button = QPushButton("Обнови")
        self.refresh_button.clicked.connect(self.load_protocols)

        self.print_button = QPushButton("Отпечатване")
        self.print_button.clicked.connect(self.print_table)

        # # Създаване на поле за търсене
        # self.search_input = QLineEdit()
        # self.search_input.textChanged.connect(self.search_devices)

        # Създаване на падащо меню за критерий за сортиране
        self.sort_criterion_combo = QComboBox()
        self.sort_criterion_combo.addItems(["ID", "Сериен номер"])  # Добавяне на възможните критерии
        self.sort_criterion_combo.currentIndexChanged.connect(
            self.sort_devices)  # Свързване на сигнала currentIndexChanged с функцията за сортиране

        # Създаване на падащо меню за посока на сортиране
        self.sort_direction_combo = QComboBox()
        self.sort_direction_combo.addItems(["Възходящо", "Низходящо"])  # Добавяне на възможните посоки
        self.sort_direction_combo.currentIndexChanged.connect(
            self.sort_devices)  # Свързване на сигнала currentIndexChanged с функцията за сортиране
        #
        # Създаване на бутон "Импортиране"
        self.import_button = QPushButton("Импортиране")
        self.import_button.clicked.connect(self.import_from_csv)

        # # Създаване на бутон "Изтрий избраните"
        # self.delete_selected_button = QPushButton("Изтрий избраните")
        # self.delete_selected_button.clicked.connect(self.delete_selected_products)

        self.serial_number_filter_input = QLineEdit()
        self.serial_number_filter_input.textChanged.connect(self.check_input)

        self.id_filter_input = QLineEdit()
        self.id_filter_input.textChanged.connect(self.check_input)

        self.date_filter_input = QLineEdit()
        self.date_filter_input.textChanged.connect(self.check_input)

        # Добавяне на полетата за филтриране към layout-а
        filter_layout = QVBoxLayout()
        filter_layout.addWidget(QLabel("ID:"))
        filter_layout.addWidget(self.id_filter_input)
        filter_layout.addWidget(QLabel("Сериен номер:"))
        filter_layout.addWidget(self.serial_number_filter_input)
        filter_layout.addWidget(QLabel("Дата:"))
        filter_layout.addWidget(self.date_filter_input)

        # Добавяне на падащите менюта към layout-а
        sort_layout = QHBoxLayout()
        sort_layout.addWidget(self.sort_criterion_combo)
        sort_layout.addWidget(self.sort_direction_combo)

        # Създаване на бутон "Експортиране"
        self.export_button = QPushButton("Експортиране")
        self.export_button.clicked.connect(self.export_to_csv)

        # Свързване на сигнала cellDoubleClicked с функцията за редактиране
        self.protocols_table.cellDoubleClicked.connect(self.open_edit_protocols_window)

        # Разположение на елементите
        button_layout = QHBoxLayout()
        # button_layout.addWidget(self.new_button)
        button_layout.addWidget(self.edit_button)
        # Добавяне на бутона "Изтрий" към layout-а
        button_layout.addWidget(self.delete_button)
        button_layout.addWidget(self.refresh_button)
        # Добавяне на бутона "Отпечатване" към layout-а
        button_layout.addWidget(self.print_button)
        button_layout.addWidget(self.export_button)
        # Добавяне на бутона "Импортиране" към layout-а
        button_layout.addWidget(self.import_button)
        # Добавяне на бутона "Изтрий избраните" към layout-а
        # button_layout.addWidget(self.delete_selected_button)

        main_layout = QVBoxLayout()
        # main_layout.addWidget(self.search_input)
        main_layout.addWidget(self.protocols_table)
        main_layout.addLayout(button_layout)
        main_layout.addLayout(sort_layout)
        main_layout.addLayout(filter_layout)

        self.setLayout(main_layout)

        self.all_protocols = []  # Създаване на празен списък за оригиналните данни

    def check_input(self):
        if not self.id_filter_input.text().strip() and not self.serial_number_filter_input.text().strip() and not self.date_filter_input.text().strip():
            self.load_protocols()
        else:
            self.load_protocols()
            self.filter_products()

    def load_protocols(self):
        # id
        # serial_number
        # problem_description
        # date
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT id, serial_number, problem_description, date FROM protocols"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            protocols = cur.fetchall()
            self.all_protocols = []  # Изчистване на списъка с оригиналните данни

            self.protocols_table.setRowCount(0)  # Изчистване на таблицата
            # self.protocols_table.insertColumn(0)
            for protocol in protocols:
                row_num = self.protocols_table.rowCount()
                self.protocols_table.insertRow(row_num)
                # checkbox = QCheckBox()
                # self.protocols_table.setCellWidget(row_num, 0, checkbox)  # Добавяне на checkbox към първата колона
                protocol_data = {}  # Речник за съхранение на данните за продукта
                for i, data in enumerate(protocol):
                    item = QTableWidgetItem(str(data))
                    self.protocols_table.setItem(row_num, i, item)  #i+1
                    protocol_data[self.column_mapping[i]] = data  # Добавяне на данните в речника

                self.all_protocols.append(protocol_data)  # Добавяне на данните за продукта към списъка с оригиналните данни

                # Задаване на подравняване вдясно за елементите от колоните с цена и количество
                item = self.protocols_table.item(row_num, 0)  # Вземане на елемента от колоната с баркода
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране
                item = self.protocols_table.item(row_num, 2)  # Вземане на елемента от колоната с цена
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

                item = self.protocols_table.item(row_num, 3)  # Вземане на елемента от колоната с количество
                item.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)  # Задаване на подравняване вдясно и вертикално центриране

            cur.close()
            conn.close()
            self.protocols_table.setColumnWidth(2, 635)  # Увеличава ширината на втората колона (индекс 1) на 200 пиксела

        except Exception as e:
            print(f"Грешка при зареждане на протоколите: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на протоколите: {e}")

    def open_new_protocols_window(self):
        self.new_protocols_window = NewProtocolWindow(self)
        result = self.new_protocols_window.exec()
        if result == QDialog.DialogCode.Accepted:
            self.load_protocols()

    def open_edit_protocols_window(self, row, column):
        selected_rows = self.protocols_table.selectedItems()
        if not selected_rows:
            QMessageBox.warning(self, "Грешка", "Моля, изберете протокол за редактиране.")
            return

        # protocol_id = selected_rows[0].text()
        # print(selected_rows[0])
        # print(selected_rows[1])
        # print(selected_rows[2])
        # print(selected_rows[3])# Вземане на ID на продукта от избрания ред
        #
        # protocol_id = self.protocols_table.item(row, 0).text() if self.protocols_table.item(row, 0) else ""
        # print(f"ID на избрания протокол: {protocol_id}")

        column_count = self.protocols_table.columnCount()
        row_data = []

        for col in range(column_count):
            item = self.protocols_table.item(row, col)  # Вземаме клетката
            row_data.append(item.text() if item is not None else "")  # Проверяваме дали е None

        print(f"Избран ред ({row}): {row_data}")  # Дебъгване

        # Ако имаш нужда от ID (първа колона)
        protocol_id = row_data[0] if row_data else ""

        print(f"ID на протокола: {protocol_id}")


        self.edit_protocols_window = EditProtocolWindow(protocol_id, self)
        result = self.edit_protocols_window.exec()
        if result == QDialog.DialogCode.Accepted:
            self.load_protocols()

    def search_devices(self, text):
        try:
            conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT id, serial_number, problem_description, date FROM protocols WHERE serial_number LIKE ?"  # Търсене по име (ILIKE е case-insensitive)
            cur.execute(sql, ('%' + text + '%',))  # Добавяне на wildcard символи (%) за търсене на подstring
            protocols = cur.fetchall()

            self.protocols_table.setRowCount(0)  # Изчистване на таблицата
            for protocol in protocols:
                row_num = self.protocols_table.rowCount()
                self.protocols_table.insertRow(row_num)
                for i, data in enumerate(protocol):
                    item = QTableWidgetItem(str(data))
                    self.protocols_table.setItem(row_num, i, item)

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при търсене на протоколите: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при търсенето на протоколите: {e}")

    def sort_devices(self):
        try:
            conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Вземане на избрания критерий и посока на сортиране
            criterion = self.sort_criterion_combo.currentText()
            direction = self.sort_direction_combo.currentText()

            # Изпълнение на SQL заявка за сортиране на продуктите
            if criterion == "":
                sql = "SELECT id, serial_number, problem_description, date FROM protocols ORDER BY id "  # SQL заявка за извличане на данните за продуктите
            elif criterion == "Сериен номер":
                sql = "SELECT id, serial_number, problem_description, date FROM protocols ORDER BY serial_number "  # SQL заявка за извличане на данните за продуктите

            if direction == "Низходящо":
                sql += "DESC"

            cur.execute(sql)
            protocols = cur.fetchall()

            self.protocols_table.setRowCount(0)  # Изчистване на таблицата
            for protocol in protocols:
                row_num = self.protocols_table.rowCount()
                self.protocols_table.insertRow(row_num)
                for i, data in enumerate(protocol):
                    item = QTableWidgetItem(str(data))
                    self.protocols_table.setItem(row_num, i, item)

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при сортиране на протоколите: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при сортирането на протоколите: {e}")

    def filter_products(self):
        try:
            # Вземане на текста от полетата за филтриране
            id_filter = str(self.id_filter_input.text()).lower()
            serial_number_filter = str(self.serial_number_filter_input.text()).lower()
            date_filter = self.date_filter_input.text()

            if id_filter == '':
                id_filter = "none"
            if serial_number_filter == '':
                serial_number_filter = "none"
            if date_filter == '':
                date_filter = "none"

            print(self.all_protocols)
            # Филтриране на данните
            filtered_protocols = []
            for protocol in self.all_protocols:
                if id_filter in str(protocol["id"]).lower():
                    filtered_protocols.append(protocol)
                elif serial_number_filter in protocol["serial_number"].lower():
                    filtered_protocols.append(protocol)
                elif date_filter in str(protocol["date"]):
                    filtered_protocols.append(protocol)
                # # Използваме копие на оригиналните данни
                # if (id_filter in protocol["company_name"].lower() or
                #         serial_number_filter in protocol["serial_number"].lower() or
                #         (date_filter in str(protocol["eik"]))):
                #     filtered_protocols.append(protocol)

            # Обновяване на таблицата
            self.protocols_table.setRowCount(0)  # Изчистване на таблицата
            for protocol in filtered_protocols:
                row_num = self.protocols_table.rowCount()
                self.protocols_table.insertRow(row_num)
                for i, (col_name, value) in enumerate(protocol.items()):
                    item = QTableWidgetItem(str(value))
                    self.protocols_table.setItem(row_num, i, item)

        except Exception as e:
            print(f"Грешка при филтриране на протоколи: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при филтрирането на протоколите: {e}")

    def print_table(self):
        try:
            # Създаване на печатащ документ
            printer = QPrinter()

            # Създаване на диалогов прозорец за настройка на печат
            print_dialog = QPrintDialog(printer, self)
            if print_dialog.exec() == QPrintDialog.DialogCode.Accepted:
                # Създаване на painter за рисуване върху печатащия документ
                painter = QPainter(printer)

                # Рисуване на данните от таблицата
                self.protocols_table.render(painter)

                painter.end()

        except Exception as e:
            print(f"Грешка при отпечатване на таблицата: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при отпечатването на таблицата: {e}")

    # def validate_data(self, row_data):
    #     try:
    #         # Проверка за празни полета
    #         if not row_data["name"]:
    #             raise ValueError("Полето 'име' е задължително.")
    #         if not row_data["barcode"]:
    #             raise ValueError("Полето 'баркод' е задължително.")
    #         if not row_data["price"]:
    #             raise ValueError("Полето 'цена' е задължително.")
    #         if not row_data["quantity"]:
    #             raise ValueError("Полето 'количество' е задължително.")
    #
    #         # Проверка за валидни типове данни
    #         float(row_data["price"])  # Проверка дали цената е число
    #         int(row_data["quantity"])  # Проверка дали количеството е цяло число
    #
    #         return True  # Данните са валидни
    #
    #     except ValueError as e:
    #         QMessageBox.warning(self, "Грешка при валидиране", str(e))
    #         return False  # Данните не са валидни

    def import_from_csv(self):
        try:
            # Отваряне на диалогов прозорец за избор на файл
            file_path, _ = QFileDialog.getOpenFileName(self, "Избор на CSV файл", "", "CSV Files (*.csv)")

            if file_path:
                # Четене на данните от CSV файла
                with open(file_path, "r", encoding="utf-8-sig") as csvfile:
                    reader = csv.DictReader(csvfile)  # Използване на DictReader за четене на данни по колони
                    data = list(reader)

                # Добавяне на данните към таблицата
                self.protocols_table.setRowCount(0)  # Изчистване на таблицата преди добавяне на нови данни
                for row_data in data:
                    row_num = self.protocols_table.rowCount()
                    self.protocols_table.insertRow(row_num)
                    for col_name, value in row_data.items():
                        col_index = list(self.column_mapping.keys())[list(self.column_mapping.values()).index(
                            col_name)]  # Намиране на индекса на колоната по име
                        item = QTableWidgetItem(value)
                        self.protocols_table.setItem(row_num, col_index, item)

                QMessageBox.information(self, "Успех", "Данните са успешно импортирани от CSV файл.")

                # Питаме потребителя дали иска да запази данните в базата данни
                reply = QMessageBox.question(self, "Запазване в база данни",
                                             "Желаете ли да запазите импортираните данни в базата данни?",
                                             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if reply == QMessageBox.StandardButton.Yes:
                    try:
                        conn = sqlite3.connect(DB_PATH) # Заменете с вашите данни за връзка
                        cur = conn.cursor()

                        # Проверяваме дали записите вече съществуват и добавяме само новите
                        for row_data in data:
                            #     if not self.validate_data(row_data):
                            #         continue  # Преминаваме към следващия запис, ако данните не са валидни
                            # Изграждаме SQL заявка за проверка дали записът вече съществува
                            check_sql = "SELECT id FROM products WHERE barcode = ?"  # Проверяваме по баркод, можете да промените критерия
                            cur.execute(check_sql, (row_data["barcode"],))
                            existing_product = cur.fetchone()

                            if existing_product is None:  # Ако записът не съществува, го добавяме
                                # Изграждаме SQL заявка за добавяне на нов запис
                                insert_sql = "INSERT INTO products (name, barcode, price, quantity) VALUES (?, ?, ?, ?)"  # Адаптирайте колоните според вашите нужди
                                values = (row_data["name"], row_data["barcode"], row_data["price"],
                                          row_data["quantity"])  # Адаптирайте данните според вашите нужди
                                cur.execute(insert_sql, values)

                        conn.commit()
                        cur.close()
                        conn.close()

                        QMessageBox.information(self, "Успех", "Данните са успешно запазени в базата данни.")
                        self.load_protocols()  # Обновяваме таблицата с продуктите след импортирането и запазването
                    except Exception as e:
                        print(f"Грешка при запазване в базата данни: {e}")
                        conn.rollback()  # Откатваме транзакцията в случай на грешка
                        QMessageBox.critical(self, "Грешка",
                                             f"Възникна грешка при запазването на данните в базата данни: {e}")


        except Exception as e:
            print(f"Грешка при импортиране от CSV файл: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при импортирането на данните от CSV файл: {e}")

    def export_to_csv(self):
        try:
            # Получаване на данните от таблицата, включително заглавките на колоните
            headers = [self.protocols_table.horizontalHeaderItem(i).text() for i in
                       range(self.protocols_table.columnCount())]
            data = []
            for row in range(self.protocols_table.rowCount()):
                row_data = {}  # Променяме на речник, за да запазим реда на колоните
                for col in range(self.protocols_table.columnCount()):
                    item = self.protocols_table.item(row, col)
                    if item is not None:
                        column_name = self.column_mapping[col]  # Вземаме името на колоната от речника
                        row_data[column_name] = item.text()
                    else:
                        column_name = self.column_mapping[col]
                        row_data[column_name] = ""  # Добавяне на празен низ, ако клетката е празна
                data.append(row_data)
                print(data)

            # Отваряне на диалогов прозорец за запазване на файл
            file_path, _ = QFileDialog.getSaveFileName(self, "Запазване на CSV файл", "", "CSV Files (*.csv)")

            if file_path:
                # Запис на данните в CSV файла с UTF-8 кодиране и BOM
                with open(file_path, "w", newline="", encoding="utf-8-sig") as csvfile:
                    fieldnames = list(self.column_mapping.values())  # Вземаме имената на колоните от речника
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=",", quoting=csv.QUOTE_NONNUMERIC)
                    writer.writeheader()  # Запис на заглавките на колоните
                    writer.writerows(data)  # Запис на данните

                QMessageBox.information(self, "Успех", "Данните са успешно експортирани в CSV файл.")
                os_name = platform.system()
                # Питаме потребителя дали иска да отвори файла
                reply = QMessageBox.question(self, "Отваряне на файл", "Искате ли да отворите файла?",
                                             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if reply == QMessageBox.StandardButton.Yes:
                    # Отваряне на файла в зависимост от операционната система
                    if os_name == "Windows":
                        os.startfile(file_path)
                    elif os_name == "Darwin":  # macOS
                        subprocess.Popen(['open', file_path])
                    elif os_name == "Linux":
                        subprocess.Popen(['xdg-open', file_path])
                    else:
                        QMessageBox.warning(self, "Грешка", "Неподдържана операционна система.")

        except Exception as e:
            print(f"Грешка при експортиране в CSV файл: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при експортирането на данните в CSV файл: {e}")

    def edit_selected_product(self, row, column):
        try:
            # Вземане на ID на избрания продукт
            protocol_id = int(self.protocols_table.item(row, 0).text())

            # Отваряне на прозореца за редактиране на продукт
            self.edit_protocols_window = EditProtocolWindow(protocol_id, self)
            result = self.edit_protocols_window.exec()

            # Обновяване на таблицата с продуктите, ако потребителят е запазил промените
            if result == QDialog.DialogCode.Accepted:
                self.load_protocols()

        except Exception as e:
            print(f"Грешка при отваряне на прозореца за редактиране: {e}")
            QMessageBox.critical(self, "Грешка",
                                 f"Възникна грешка при отварянето на прозореца за редактиране: {e}")

    def delete_product(self):
        selected_items = self.protocols_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Грешка", "Моля, изберете протокол за изтриване.")
            return

        protocol_id = selected_items[0].text()

        reply = QMessageBox.question(self, "Потвърждение", "Сигурни ли сте, че искате да изтриете този протокол?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            try:
                conn = sqlite3.connect(DB_PATH)# Заменете с вашите данни за връзка
                cur = conn.cursor()

                sql = "DELETE FROM protocols WHERE id = ?"
                cur.execute(sql, (protocol_id,))

                conn.commit()
                cur.close()
                conn.close()

                self.load_protocols()  # Обновяване на таблицата с продуктите

                QMessageBox.information(self, "Успех", "Протокола е успешно изтрит.")

            except Exception as e:
                print(f"Грешка при изтриване на протокол: {e}")
                QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на протокола: {e}")

    def delete_selected_products(self):
        try:
            # Вземане на избраните продукти
            selected_products = []
            for row in range(self.protocols_table.rowCount()):
                checkbox = self.protocols_table.cellWidget(row, 0)
                if checkbox.isChecked():
                    product_id = int(self.protocols_table.item(row, 1).text())  # ID-то е във втората колона (индекс 1)
                    selected_products.append(product_id)

            if not selected_products:
                QMessageBox.warning(self, "Грешка", "Моля, изберете продукти за изтриване.")
                return

            # Показване на диалогов прозорец за потвърждение на изтриването
            reply = QMessageBox.question(self, "Потвърждение",
                                         "Сигурни ли сте, че искате да изтриете избраните продукти?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

            if reply == QMessageBox.StandardButton.Yes:
                try:
                    conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
                    cur = conn.cursor()

                    # Изтриване на избраните продукти от базата данни
                    for product_id in selected_products:
                        sql = "DELETE FROM products WHERE id = ?"
                        cur.execute(sql, (product_id,))

                    conn.commit()
                    cur.close()
                    conn.close()

                    self.load_protocols()  # Обновяване на таблицата с продуктите

                    QMessageBox.information(self, "Успех", "Избраните продукти са успешно изтрити.")

                except Exception as e:
                    print(f"Грешка при изтриване на продукти: {e}")
                    QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на продуктите: {e}")

        except Exception as e:
            print(f"Грешка при изтриване на продукти: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при изтриването на продуктите: {e}")

#window for new protocol
class NewProtocolWindow(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        #["ID", "Сериен номер", "Описание на проблема" "Дата"])  # Задаване на заглавия на колоните

        self.setWindowTitle("Нов протокол")
        self.all_ids = []  # Изчистване на списъка с оригиналните данни

        self.load_protocols()
        if len(self.all_ids) > 0:
            id_nums = len(self.all_ids)
            new_id_num = id_nums + 1
            id_text = str(new_id_num)
        else:
            id_text = "1"

        # Създаване на полета за въвеждане на данни
        self.id_label = QLabel("ID:")
        self.id_input = QLineEdit().setText(id_text)

        self.serial_number_label = QLabel("Серйен номер:")
        self.serial_number_input = QLineEdit()

        self.problem_description_label = QLabel("Описание на проблема:")
        self.problem_description_input = QTextEdit()

        # Създаване на полета за въвеждане на описание
        self.date_label = QLabel("Дата:")
        self.date_input = QDateEdit()
        self.date_input.setCalendarPopup(True)
        self.date_input.setDisplayFormat("dd.MM.yyyy")
        self.date_input.setDate(QDate.currentDate())


        # Създаване на бутон за запазване
        self.save_button = QPushButton("Запази")
        self.save_button.clicked.connect(self.save_protocol)

        # Създаване на бутон за отказ
        self.cancel_button = QPushButton("Отказ")
        self.cancel_button.clicked.connect(self.close)

        # Разположение на елементите в прозореца
        layout = QVBoxLayout()
        layout.addWidget(self.id_label)
        layout.addWidget(self.id_input)
        layout.addWidget(self.serial_number_label)
        layout.addWidget(self.serial_number_input)
        layout.addWidget(self.problem_description_label)
        layout.addWidget(self.problem_description_input)
        layout.addWidget(self.date_label)
        layout.addWidget(self.date_input)
        layout.addWidget(self.save_button)
        layout.addWidget(self.cancel_button)
        self.setLayout(layout)

    # def upload_image(self):
    #     try:
    #         # Отваряне на диалогов прозорец за избиране на файл
    #         # options = QFileDialog.Options()
    #         file_path, _ = QFileDialog.getOpenFileName(self, "Изберете снимка", "", "Image Files (*.png *.jpg *.jpeg)")
    #
    #         if file_path:
    #             # Зареждане на снимката и показване в QLabel
    #             pixmap = QPixmap(file_path)
    #             pixmap = pixmap.scaled(100, 100)  # Мащабиране на снимката до размера на QLabel
    #             self.image_label.setPixmap(pixmap)
    #             self.image_path = file_path  # Запазване на пътя към снимката
    #     except Exception as e:
    #         print(f"Грешка при качване на снимка: {e}")
    #         QMessageBox.critical(self, "Грешка", f"Възникна грешка при качването на снимката: {e}")

    def save_protocol(self):
        # Тук ще добавим код за запазване на продукта в базата данни
        # Тук ще добавим код за запазване на продукта в базата данни
        id = self.id_input.text()
        serial_number = self.serial_number_input.text()
        problem_description = self.problem_description_input.text()
        date = self.date_input.text()


        if not id or not serial_number or not problem_description or not date:
            QMessageBox.warning(self, "Грешка", "Моля, попълнете всички полета.")
            return

        try:
            # Свързване с базата данни
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за вмъкване на нов продукт
            sql = "INSERT INTO protocols (id, serial_number, problem_description, date) VALUES (?, ?, ?, ?)"
            values = (id, serial_number, problem_description, date)
            cur.execute(sql, values)

            # Запис на промените в базата данни
            conn.commit()

            # Затваряне на връзката с базата данни
            cur.close()
            conn.close()

            QMessageBox.information(self, "Успех", "Протоколът е успешно добавено.")
            self.accept()  # Затваря прозореца след успешно запазване

        except Exception as e:
            print(f"Грешка при запазване на устройството: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при запазването на протоколът: {e}")

    def load_protocols(self):
        # id
        # serial_number
        # problem_description
        # date
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            sql = "SELECT id FROM protocols"  # SQL заявка за извличане на данните за продуктите
            cur.execute(sql)
            ids = cur.fetchall()
            self.all_ids = []  # Изчистване на списъка с оригиналните данни

            # self.protocols_table.insertColumn(0)
            for id in ids:
                self.all_ids.append(id)  # Добавяне на данните за продукта към списъка с оригиналните данни
            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на ид-тата: {e}")

#window for edit protocol
class EditProtocolWindow(QDialog):
    def __init__(self, protocol_id, parent=None):
        super().__init__(parent)

        self.setWindowTitle("Редактиране на протоколи")
        self.protocol_id = protocol_id  # Запазване на ID на продукта

        # Създаване на полета за въвеждане на данни
        self.id_label = QLabel("ID:")
        self.id_input = QLineEdit()

        self.serial_number_label = QLabel("Серйен номер:")
        self.serial_number_input = QLineEdit()

        self.problem_description_label = QLabel("Описание на проблема:")
        self.problem_description_input = QTextEdit()

        # Създаване на полета за въвеждане на описание
        # Създаване на полета за въвеждане на описание
        self.date_label = QLabel("Дата:")
        self.date_input = QDateEdit()
        self.date_input.setCalendarPopup(True)  # Позволява изскачащ календар

        # Създаване на бутон за запазване
        self.save_button = QPushButton("Запази")
        self.save_button.clicked.connect(self.save_protocol)

        # Създаване на бутон за отказ
        self.cancel_button = QPushButton("Отказ")
        self.cancel_button.clicked.connect(self.close)

        # Разположение на елементите в прозореца
        layout = QVBoxLayout()
        layout.addWidget(self.id_label)
        layout.addWidget(self.id_input)
        layout.addWidget(self.serial_number_label)
        layout.addWidget(self.serial_number_input)
        layout.addWidget(self.problem_description_label)
        layout.addWidget(self.problem_description_input)
        layout.addWidget(self.date_label)
        layout.addWidget(self.date_input)
        layout.addWidget(self.save_button)
        layout.addWidget(self.cancel_button)
        self.setLayout(layout)

        # Зареждане на информацията за продукта от базата данни
        self.load_protocol_data()

    def load_protocol_data(self):
        try:
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за извличане на информация за продукта
            sql = "SELECT id, serial_number, problem_description, date FROM protocols WHERE id = ?"
            cur.execute(sql, (self.protocol_id,))
            protocol_data = cur.fetchone()

            if protocol_data:
                id, serial_number, problem_description, date = protocol_data

                self.id_input.setText(str(id))
                self.serial_number_input.setText(serial_number)
                self.problem_description_input.setText(problem_description)
                self.date_input.setDate(QDate.fromString(date, "dd.MM.yyyy"))


            else:
                QMessageBox.warning(self, "Грешка", "Протокола не е намерен.")
                self.close()

            cur.close()
            conn.close()

        except Exception as e:
            print(f"Грешка при зареждане на данни за протокола: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при зареждането на данните за протокола: {e}")

    # def upload_image(self):
    #     try:
    #         # Отваряне на диалогов прозорец за избиране на файл
    #         # options = QFileDialog.Options()
    #         file_path, _ = QFileDialog.getOpenFileName(self, "Изберете снимка", "", "Image Files (*.png *.jpg *.jpeg)")
    #
    #         if file_path:
    #             # Зареждане на снимката и показване в QLabel
    #             pixmap = QPixmap(file_path)
    #             pixmap = pixmap.scaled(100, 100)  # Мащабиране на снимката до размера на QLabel
    #             self.image_label.setPixmap(pixmap)
    #             self.image_path = file_path  # Запазване на пътя към снимката
    #     except Exception as e:
    #         print(f"Грешка при качване на снимка: {e}")
    #         QMessageBox.critical(self, "Грешка", f"Възникна грешка при качването на снимката: {e}")

    def save_protocol(self):
        # Тук ще добавим код за запазване на продукта в базата данни
        id = self.id_input.text()
        serial_number = self.serial_number_input.text()
        problem_description = self.problem_description_input.text()
        date = self.date_input.text()


        try:
            # Свързване с базата данни
            conn = sqlite3.connect(DB_PATH)  # Заменете с вашите данни за връзка
            # Заменете с вашите данни за връзка
            cur = conn.cursor()

            # Изпълнение на SQL заявка за актуализиране на информацията за продукта
            sql = "UPDATE products SET serial_number = ?, problem_description = ?, date = ? WHERE id = ?"
            values = (serial_number, problem_description, date, self.device_id)
            cur.execute(sql, values)

            conn.commit()
            cur.close()
            conn.close()

            QMessageBox.information(self, "Успех", "Протокола е успешно редактирано.")
            self.accept()

        except Exception as e:
            print(f"Грешка при редактиране на протокол: {e}")
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при редактирането на протокола: {e}")


    # def generate_protocol(self):
    #     instance_main = MainWindow()
    #
    #     if self.device_data:
    #         company_name, company_address, company_manager, device_address, phone_number, device, self.serial_number, eik, fdrid = self.device_data
    #
    #     document = Document("template.docx")
    #
    #     protocol_id = self.problems[0]
    #     date_string = self.problems[2] # Форматиране на датата
    #     # company_name = self.device_data[0]
    #     # company_address = self.device_info[1]
    #     # company_manager = self.device_info[2]
    #     # device_address = self.device_info[3]
    #     # phone_number = self.device_info[4]
    #     # device = self.device_info[5]
    #     # serial_number = self.device_data[6]
    #     problem_description = self.problems[1]
    #     first_row = protocol_id + '/' + date_string
    #
    #     for paragraph in document.paragraphs:
    #         print(paragraph.text)
    #         if "[номер на протокол от базата данни]/[дата[дд-мм-гг]]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[номер на протокол от базата данни]/[дата[дд-мм-гг]]", first_row)
    #         if "[Име на фирма]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[Име на фирма]", company_name)
    #         if "[адрес на фирма]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[адрес на фирма]", company_address)
    #         if "[управител]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[управител]", company_manager)
    #         if "[адрес на устройството]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[адрес на устройството]", device_address)
    #         if "[телефонен номер]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[телефонен номер]", phone_number)
    #         if "[какво е оставено и име и модел]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[какво е оставено и име и модел]", device)
    #         if "[сериен номер]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[сериен номер]", self.serial_number)
    #         if "[описание на порблема]" in paragraph.text:
    #             paragraph.text = paragraph.text.replace("[описание на порблема]", problem_description)
    #         # elif "[Сериен номер]" in paragraph.text:
    #         #     paragraph.text = paragraph.text.replace("[Сериен номер]", serial_number)
    #         # elif "[Модел]" in paragraph.text:
    #         #     paragraph.text = paragraph.text.replace("[Модел]", device[1])
    #         # elif "[Производител]" in paragraph.text:
    #         #     paragraph.text = paragraph.text.replace("[Производител]", device[2])
    #         # elif "[Описание на проблема]" in paragraph.text:
    #         #     paragraph.text = paragraph.text.replace("[Описание на проблема]", problem_description)
    #     pre_file_path = f"protocol_{protocol_id}_{self.serial_number}_{date_string}.docx"
    #
    #     file_path, _ = QFileDialog.getSaveFileName(self, "Запазване на Docx файл", os.path.join(os.path.expanduser("~"), "Documents", pre_file_path), "Docx Files (*.docx);; All Files (*)")
    #     file_path = file_path+''+pre_file_path
    #     document.save(f"{file_path}")
    #
    #     QMessageBox.information(self, "Успех", "Протоколът беше генериран успешно.")
    #     self.remove_fields_auto()
    #     os_name = platform.system()
    #     # Питаме потребителя дали иска да отвори файла
    #     reply = QMessageBox.question(self, "Отваряне на файл", "Искате ли да отворите файла?",
    #                                  QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
    #
    #     if reply == QMessageBox.StandardButton.Yes:
    #         # Отваряне на файла в зависимост от операционната система
    #         if os_name == "Windows":
    #             os.startfile(file_path)
    #         elif os_name == "Darwin":  # macOS
    #             subprocess.Popen(['open', file_path])
    #         elif os_name == "Linux":
    #             subprocess.Popen(['xdg-open', file_path])
    #         else:
    #             QMessageBox.warning(self, "Грешка", "Неподдържана операционна система.")



if __name__ == "__main__":
    app = QApplication(sys.argv)
    # Изпълнява се при стартиране
    initialize_database()
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
